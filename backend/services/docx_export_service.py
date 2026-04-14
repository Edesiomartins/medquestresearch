from __future__ import annotations

from io import BytesIO
from typing import Dict

from docx import Document  # type: ignore[reportMissingImports]


def build_meta_docx(meta_result: Dict) -> bytes:
    document = Document()
    document.add_heading("Meta-análise Científica", level=1)

    question = meta_result.get("question") or "Pergunta não informada"
    document.add_paragraph(f"Questão da revisão: {question}")

    pooled = meta_result.get("pooled_result") or {}
    heterogeneity = meta_result.get("heterogeneity") or {}
    bias = meta_result.get("publication_bias") or {}
    prisma = meta_result.get("prisma") or {}

    document.add_heading("Resumo Quantitativo", level=2)
    if pooled:
        document.add_paragraph(
            (
                f"Modelo: {pooled.get('model')} | Medida: {pooled.get('effect_measure')} | "
                f"Efeito combinado: {pooled.get('effect')} "
                f"(IC95%: {pooled.get('ci_low')} a {pooled.get('ci_high')}) | "
                f"p global: {pooled.get('p_value')}"
            )
        )
    else:
        document.add_paragraph("Pooling quantitativo não disponível.")

    if heterogeneity:
        document.add_paragraph(
            (
                f"Heterogeneidade — Q: {heterogeneity.get('Q')}, df: {heterogeneity.get('df')}, "
                f"I²: {heterogeneity.get('I2')}%, tau²: {heterogeneity.get('tau2')}, "
                f"p heterogeneidade: {heterogeneity.get('p_heterogeneity')}"
            )
        )

    if bias:
        document.add_paragraph(
            (
                f"Viés de publicação — Egger p: {bias.get('egger_p')}, "
                f"Begg p: {bias.get('begg_p')}. Interpretação: {bias.get('interpretation')}"
            )
        )

    document.add_heading("Fluxo PRISMA", level=2)
    document.add_paragraph(
        (
            f"Identificados: {prisma.get('identified', 0)} | "
            f"Triados: {prisma.get('screened', 0)} | "
            f"Elegíveis: {prisma.get('eligible', 0)} | "
            f"Incluídos: {prisma.get('included', 0)}"
        )
    )
    excluded = prisma.get("excluded_reasons") or []
    if excluded:
        document.add_paragraph("Motivos de exclusão:")
        for item in excluded:
            document.add_paragraph(
                f"- {item.get('study_id', 'study')}: {item.get('reason', 'sem motivo')}",
                style="List Bullet",
            )

    narrative = meta_result.get("narrative_results")
    if narrative:
        document.add_heading("Síntese Narrativa", level=2)
        document.add_paragraph(narrative)

    sections = meta_result.get("article_sections") or {}
    if sections:
        document.add_heading("Seções do Manuscrito", level=2)
        for section_name in [
            "abstract",
            "introduction",
            "methods",
            "results",
            "discussion",
            "conclusion",
        ]:
            text = sections.get(section_name)
            if not text:
                continue
            document.add_heading(section_name.capitalize(), level=3)
            document.add_paragraph(str(text))

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()

