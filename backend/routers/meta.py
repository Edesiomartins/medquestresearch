from __future__ import annotations

import os
import tempfile
import time
from typing import List

from fastapi import APIRouter, Depends, File, Header, HTTPException, UploadFile
from fastapi.responses import StreamingResponse

from backend.schemas.article import ArticleSectionResponse
from backend.schemas.meta import (
    MetaAnalysisResponse,
    MetaAnalyzeRequest,
    MetaReviewRequest,
    MetaUploadResponse,
    MetaPipelineStatus,
)
from backend.services.article_generation_service import (
    VALID_SECTIONS,
    generate_article_section,
)
from backend.services.docx_export_service import build_meta_docx
from backend.services.export_bundle_service import build_submission_zip
from backend.services.extraction_service import extract_studies_from_texts
from backend.services.meta_pipeline_service import analyze_meta
from backend.services.review_validation_service import validate_reviewed_studies

try:
    from backend.database import db_select_one
except Exception:
    from database import db_select_one  # type: ignore

try:
    from backend.pdf_processor import extrair_texto_pdf
except Exception:
    from pdf_processor import extrair_texto_pdf  # type: ignore


router = APIRouter(prefix="/api/meta", tags=["meta-analysis-v2"])

TITLE_MARKER = "[[MEDQUEST_TITLE]]:"


def _extract_docx_text(path: str, filename: str) -> str:
    from docx import Document  # python-docx

    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    title = os.path.splitext(os.path.basename(filename))[0].replace("_", " ").strip()
    body = "\n".join(parts)
    return f"{TITLE_MARKER} {title}\n{body}" if title else body


def require_api_key(authorization: str = Header(None)):
    if not authorization:
        raise HTTPException(status_code=401, detail="Não autorizado")
    token = authorization.replace("Bearer ", "").strip()
    user = db_select_one("SELECT * FROM usuarios WHERE token = %s", (token,))
    if not user:
        raise HTTPException(status_code=401, detail="Não autorizado")
    return dict(user)


@router.post("/upload", response_model=MetaUploadResponse)
async def upload_studies(files: List[UploadFile] = File(...), user=Depends(require_api_key)):
    if len(files) == 0:
        raise HTTPException(status_code=400, detail="Nenhum arquivo enviado.")
    if len(files) > 25:
        raise HTTPException(status_code=400, detail="Máximo de 25 arquivos por upload.")

    texts: List[str] = []
    notes: List[str] = []
    failed_files: List[str] = []
    for file in files:
        if not file.filename or not file.filename.lower().endswith((".pdf", ".docx")):
            raise HTTPException(status_code=400, detail=f"Formato não suportado: {file.filename}")
        payload = await file.read()
        if len(payload) > 20 * 1024 * 1024:
            raise HTTPException(
                status_code=400,
                detail=f"Arquivo {file.filename} excede 20MB.",
            )
        safe_suffix = ".pdf" if file.filename.lower().endswith(".pdf") else ".docx"
        temp_name = os.path.join(
            tempfile.gettempdir(),
            f"meta_upload_{time.time_ns()}{safe_suffix}",
        )
        with open(temp_name, "wb") as stream:
            stream.write(payload)
        try:
            if safe_suffix == ".pdf":
                text = extrair_texto_pdf(temp_name)
                if isinstance(text, list):
                    text = "\n".join(text)
            else:
                text = _extract_docx_text(temp_name, file.filename)
            if not (text or "").strip():
                failed_files.append(file.filename)
                notes.append(f"{file.filename}: nenhum texto extraível (PDF escaneado sem OCR?). Arquivo ignorado.")
                continue
            texts.append(text)
        except Exception as error:
            failed_files.append(file.filename)
            notes.append(f"{file.filename}: falha ao processar ({type(error).__name__}). Arquivo ignorado.")
        finally:
            if os.path.exists(temp_name):
                os.remove(temp_name)

    if not texts:
        raise HTTPException(
            status_code=422,
            detail="Nenhum arquivo pôde ser processado. " + " ".join(notes),
        )

    project_id = f"meta_{time.time_ns()}"
    studies = extract_studies_from_texts(project_id=project_id, texts=texts, notes=notes)
    status = MetaPipelineStatus.warning if failed_files else MetaPipelineStatus.success
    notes.insert(0, f"{len(texts)} de {len(files)} arquivos processados.")
    return MetaUploadResponse(
        status=status,
        project_id=project_id,
        studies=studies,
        notes=notes,
    )


@router.post("/extract", response_model=MetaUploadResponse)
def extract_structured(payload: MetaReviewRequest, user=Depends(require_api_key)):
    _ = user
    project_id = payload.project_id or str(int(time.time()))
    return MetaUploadResponse(
        status=MetaPipelineStatus.success,
        project_id=project_id,
        studies=payload.studies,
        notes=["Extração estruturada pronta para revisão humana."],
    )


@router.post("/review", response_model=MetaUploadResponse)
def review_extraction(payload: MetaReviewRequest, user=Depends(require_api_key)):
    _ = user
    project_id = payload.project_id or str(int(time.time()))
    reviewed, notes = validate_reviewed_studies(payload.studies)
    return MetaUploadResponse(
        status=MetaPipelineStatus.success,
        project_id=project_id,
        studies=reviewed,
        notes=notes,
    )


@router.post("/analyze", response_model=MetaAnalysisResponse)
def analyze(payload: MetaAnalyzeRequest, user=Depends(require_api_key)):
    _ = user
    if len(payload.question or "") > 1000:
        raise HTTPException(status_code=400, detail="Pergunta de pesquisa muito longa (máx. 1000 caracteres).")
    return analyze_meta(payload)


@router.post("/plots", response_model=dict)
def plots(payload: MetaAnalyzeRequest, user=Depends(require_api_key)):
    _ = user
    result = analyze_meta(payload)
    return {
        "status": result.status,
        "project_id": result.project_id,
        "forest_plot_svg": result.forest_plot_svg,
        "funnel_plot_svg": result.funnel_plot_svg,
    }


@router.post("/article", response_model=ArticleSectionResponse)
def article(payload: MetaAnalyzeRequest, section: str = "results", user=Depends(require_api_key)):
    _ = user
    if section not in VALID_SECTIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Seção inválida: {section}. Use uma de: {', '.join(sorted(VALID_SECTIONS))}.",
        )
    # A seção solicitada é gerada à parte; evita gerar as 6 seções dentro da análise.
    payload.generate_article = False
    result = analyze_meta(payload)
    content = generate_article_section(result.dict(), section)
    if not content:
        raise HTTPException(status_code=502, detail=f"Falha ao gerar a seção '{section}'. Tente novamente.")
    return ArticleSectionResponse(
        section=section,
        content=content,
        warnings=result.warnings,
    )


@router.post("/export/docx")
def export_docx(
    payload: MetaAnalyzeRequest,
    filename: str = "meta_analise",
    user=Depends(require_api_key),
):
    _ = user
    result = analyze_meta(payload)
    docx_bytes = build_meta_docx(result.dict())
    safe_name = "".join(ch if ch.isalnum() or ch in ("_", "-") else "_" for ch in filename).strip("_") or "meta_analise"
    headers = {
        "Content-Disposition": f'attachment; filename="{safe_name}.docx"'
    }
    return StreamingResponse(
        iter([docx_bytes]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@router.post("/export/zip")
def export_zip(
    payload: MetaAnalyzeRequest,
    filename: str = "meta_analise_submission",
    user=Depends(require_api_key),
):
    _ = user
    result = analyze_meta(payload)
    zip_bytes = build_submission_zip(result.dict())
    safe_name = "".join(ch if ch.isalnum() or ch in ("_", "-") else "_" for ch in filename).strip("_") or "meta_analise_submission"
    headers = {
        "Content-Disposition": f'attachment; filename="{safe_name}.zip"'
    }
    return StreamingResponse(
        iter([zip_bytes]),
        media_type="application/zip",
        headers=headers,
    )

